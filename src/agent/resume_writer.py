"""Generates a tailored resume for one job by editing a COPY of your actual
base resume .docx — not building a new document from scratch.

Earlier version regenerated the whole resume as plain text via an LLM and
rebuilt a docx with guessed styling — it never looked like your real resume
because it wasn't your real resume, just a lookalike. This version copies
your actual file (so every existing style, font, tab stop, bullet formatting
is byte-for-byte preserved) and only does two surgical edits:
  1. Rewrites the PROFESSIONAL SUMMARY paragraph's text in place (same run,
     same formatting, just different words).
  2. Inserts 1-3 new bullet points, cloned from an existing bullet paragraph's
     exact XML (so they're formatted identically to the ones already there),
     into whichever existing job entries are most relevant.
Nothing else in the document is touched — same employers, same dates, same
other bullets, same layout. "Add points to it," not "regenerate it."

Never invents employers, titles, dates, or technologies that aren't already
established elsewhere in the resume — the prompt is explicit about that.
"""

import copy
import re
from pathlib import Path
from uuid import UUID

import docx
from docx.text.paragraph import Paragraph
from pydantic import BaseModel, Field

from langchain_openai import ChatOpenAI

from storage.jobs import get_job
from storage.resume import get_base_resume
from storage.resumes import insert_resume

_MODEL = "gpt-4o-mini"
OUTPUT_DIR = Path(__file__).resolve().parent.parent.parent / "output" / "resumes"

_JOB_TITLE_RE = re.compile(r"\|.*(—|--)")  # "Title  |  Company  —  Location\tDates"


class NewBullet(BaseModel):
    after_paragraph_index: int = Field(
        description="The paragraph index, taken from the numbered bullet list provided, "
        "immediately after which to insert this new bullet"
    )
    text: str = Field(description="The new bullet's text, without a leading bullet character")


class ResumeEdit(BaseModel):
    summary_replacement: str = Field(
        description=(
            "Rewritten PROFESSIONAL SUMMARY paragraph (2-4 sentences). Must only reframe "
            "emphasis using skills/technologies/experience that already appear elsewhere in "
            "the resume — do not introduce anything not already established."
        )
    )
    new_bullets: list[NewBullet] = Field(
        description=(
            "1-3 new bullet points to ADD to the most relevant existing job entries. Each must "
            "describe real experience already evidenced elsewhere in the resume (e.g. combining "
            "or re-surfacing facts already present), phrased to speak directly to the target "
            "job's requirements. Never invent a technology, employer, or achievement not already "
            "in the resume."
        )
    )


_editor = ChatOpenAI(model=_MODEL, temperature=0.2).with_structured_output(ResumeEdit)


def _slugify(text: str) -> str:
    return re.sub(r"[^a-zA-Z0-9]+", "_", text).strip("_")[:40]


def _extract_bullets(doc: docx.Document) -> list[tuple[int, str, str]]:
    """(paragraph_index, job_context, bullet_text) for every bullet in the resume."""
    bullets = []
    current_job = None
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        if text.startswith("•"):
            if current_job:
                bullets.append((i, current_job, text.lstrip("•").strip()))
        elif _JOB_TITLE_RE.search(text):
            current_job = text
    return bullets


def _find_summary_paragraph(doc: docx.Document) -> int | None:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "PROFESSIONAL SUMMARY":
            # the summary body is the next non-empty paragraph
            for j in range(i + 1, len(doc.paragraphs)):
                if doc.paragraphs[j].text.strip():
                    return j
    return None


def _insert_bullet_after(doc: docx.Document, after_index: int, text: str) -> None:
    """Clone the paragraph at after_index (must be an existing bullet) and set
    the clone's text — this copies its exact run formatting (font, size,
    color), so the new bullet is indistinguishable in style from the old."""
    reference = doc.paragraphs[after_index]
    new_p_element = copy.deepcopy(reference._p)
    reference._p.addnext(new_p_element)
    new_para = Paragraph(new_p_element, reference._parent)
    # collapse to a single run carrying the reference's formatting
    new_para.runs[0].text = f"•  {text}"
    for run in new_para.runs[1:]:
        run.text = ""


def generate_tailored_resume(job_id: UUID) -> dict:
    job = get_job(job_id)
    if job is None:
        raise ValueError(f"No job found with id {job_id}")

    resume = get_base_resume()
    if resume is None:
        raise ValueError("No base resume found — run resume_ingest.py first")

    doc = docx.Document(resume["file_path"])
    bullets = _extract_bullets(doc)
    summary_index = _find_summary_paragraph(doc)

    bullets_listing = "\n".join(f"[{i}] ({job_ctx}) {text}" for i, job_ctx, text in bullets)

    prompt = f"""You are tailoring a candidate's resume for a specific job by adding to it,
not rewriting it. You will NOT see or produce the full resume text — only two
things: a replacement for the professional summary, and new bullet points to
insert into existing job entries.

CURRENT PROFESSIONAL SUMMARY:
{doc.paragraphs[summary_index].text if summary_index is not None else "(not found)"}

EXISTING BULLET POINTS (index, which job, text):
{bullets_listing}

TARGET JOB:
Title: {job['title']}
Company: {job['company']}

Description:
{job['description']}

Rewrite the summary and propose 1-3 new bullets using ONLY facts, skills, and
technologies already evidenced in the bullets above. Pick after_paragraph_index
values only from the indices shown in brackets."""

    edit = _editor.invoke(prompt)

    valid_indices = {i for i, _, _ in bullets}
    if summary_index is not None:
        doc.paragraphs[summary_index].runs[0].text = edit.summary_replacement
        for run in doc.paragraphs[summary_index].runs[1:]:
            run.text = ""

    applied_bullets = []
    # insert from the bottom up so earlier indices don't shift as we add paragraphs
    for bullet in sorted(edit.new_bullets, key=lambda b: b.after_paragraph_index, reverse=True):
        if bullet.after_paragraph_index not in valid_indices:
            continue
        _insert_bullet_after(doc, bullet.after_paragraph_index, bullet.text)
        applied_bullets.append(bullet.text)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    filename = f"{_slugify(job['company'])}_{_slugify(job['title'])}_{str(job_id)[:8]}.docx"
    out_path = OUTPUT_DIR / filename
    doc.save(out_path)

    summary_of_changes = (
        f"Rewrote professional summary to emphasize fit for this role. "
        f"Added {len(applied_bullets)} new bullet(s): " + "; ".join(applied_bullets)
        if applied_bullets
        else "Rewrote professional summary to emphasize fit for this role."
    )
    resume_id = insert_resume(job_id, str(out_path), summary_of_changes)
    return {
        "resume_id": resume_id,
        "file_path": str(out_path),
        "summary_of_changes": summary_of_changes,
        "new_summary": edit.summary_replacement,
        "new_bullets": applied_bullets,
    }
